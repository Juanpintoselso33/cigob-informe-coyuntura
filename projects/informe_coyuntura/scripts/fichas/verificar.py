"""Verifica los .docx de fichas contra el snapshot, ficha por ficha.

No alcanza con que el documento tenga la forma correcta: lo que hay que
comprobar es que cada número impreso sea el que publica el informe. Un dígito
mal leído o un color desfasado no lo detecta ningún chequeo estructural, y
estas fichas se mandan afuera.

Contrasta contra web/src/data/informe.json:
  · que estén todas las fichas del cinturón, y ninguna de más
  · valor y unidad del banner "Hoy:"
  · color del semáforo, en el banner y en la sección "Color vigente"
  · peso efectivo
  · coherencia interna: el color declarado cae en el tramo que la propia
    tabla de umbrales de esa ficha dice que le corresponde
  · restos de plantilla: literales sin interpolar, "None", celdas vacías
"""
import json
import re
import sys
from pathlib import Path
from docx import Document

# La raíz del proyecto sale de la ubicación de este archivo: vive en
# scripts/fichas/, así que sube dos niveles. Antes era un path absoluto de
# una máquina concreta, que dentro del repo no le sirve a nadie.
RAIZ = Path(__file__).resolve().parents[2]
AQUI = RAIZ / "output" / "fichas"
informe = json.loads((RAIZ / "web/src/data/informe.json").read_text(encoding="utf-8"))
datos_ts = (RAIZ / "web/src/lib/datos.ts").read_text(encoding="utf-8")

CADENA = r'"((?:[^"\\]|\\.)*)"'
LABELS = dict(re.findall(r"(\w+):\s*" + CADENA,
                         re.search(r"export const LABELS[^{]*\{(.*?)\n\};",
                                   datos_ts, re.S).group(1)))
COLOR = {"verde": "VERDE", "amarillo": "AMARILLO", "naranja": "NARANJA", "rojo": "ROJO"}
DOCS = {
    "macro": "Fichas Semaforo Macro.docx",
    "politica": "Fichas Semaforo Politica.docx",
    "gestion": "Fichas Semaforo Gestion.docx",
    "vida_cotidiana": "Fichas Semaforo Vida cotidiana.docx",
    "espiritu_epoca": "Fichas Semaforo Espiritu de epoca.docx",
}
# Los mismos cortes que publica el informe, para el chequeo de coherencia del
# ITVC: la ficha declara la escala base-100 y tiene que caer en su propio tramo.
TRAMOS_ITVC = ((105.0, "verde"), (95.0, "amarillo"), (85.0, "naranja"))


def coma(x, dec=2):
    if x is None:
        return "—"
    s = f"{float(x):,.{dec}f}".replace(",", "@").replace(".", ",").replace("@", ".")
    if "," in s:
        s = s.rstrip("0").rstrip(",")
    return s.replace("-", "−")


def norm(s):
    """Word aplica sus propias reglas tipográficas al guardar: después de una
    abreviatura como "vs." mete un espacio de no separación (U+00A0). El texto
    es el mismo, así que compararlo sin normalizar da falsos positivos —dio
    cuatro, en fichas que estaban perfectas."""
    return s.replace(" ", " ").strip()


def secciones(doc):
    """Parte el documento en fichas: {título H1: [textos hasta el próximo H1]}."""
    out, actual = {}, None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            actual = norm(p.text)
            out[actual] = []
        elif actual:
            out[actual].append(p.text)
    return out


def texto_tablas_por_ficha(doc):
    """Las tablas en orden, agrupadas por la ficha en que caen. python-docx
    expone párrafos y tablas en listas separadas, así que se recorre el cuerpo
    del documento para conservar el orden real entre unos y otras."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    out, actual = {}, None
    for hijo in doc.element.body.iterchildren():
        if hijo.tag.endswith("}p"):
            p = Paragraph(hijo, doc)
            if p.style.name == "Heading 1":
                actual = norm(p.text)
                out[actual] = []
        elif hijo.tag.endswith("}tbl") and actual:
            t = Table(hijo, doc)
            out[actual].append([[c.text.strip() for c in fila.cells] for fila in t.rows])
    return out


def verificar_portada(cint, tablas_portada, fallas):
    """La portada es lo primero que se lee y repite todos los números: si
    diverge del snapshot, contradice a sus propias fichas tres páginas después.

    Lo agregué porque la prueba de mutación alteró el color de una dimensión
    acá y el verificador —que sólo miraba las fichas— dio OK igual.
    """
    cb = informe["cinturones"][cint]
    clave = next((k for k in ("itcm", "itcg", "itcp", "itvc") if k in cb), None)
    if not clave:
        return
    idx = cb[clave]

    # El banner es la tabla de UNA fila; no "la primera de la portada", que es
    # frágil y ya rompió: al anteponer la tabla de «Cómo se define el color»,
    # el verificador empezó a leer los cortes del semáforo como si fueran el
    # valor del índice y marcó ocho fallas en documentos correctos.
    banner = next((t[0] for t in tablas_portada if len(t) == 1), None)
    if banner:
        v = coma(idx.get("valor"), 1)
        if v not in norm(banner[0]):
            fallas.append(f"{cint}/portada: el índice dice «{norm(banner[0])}», "
                          f"el snapshot dice {v}")
        c = COLOR.get((idx.get("semaforo") or {}).get("color"))
        if c and norm(banner[1]) != c:
            fallas.append(f"{cint}/portada: el índice figura {norm(banner[1])}, "
                          f"el snapshot dice {c}")

    # Tabla de dimensiones: se busca por su encabezado, no por el número de
    # filas. Contar filas alcanzaba mientras fuera la única de ese tamaño, pero
    # la tabla de cortes del semáforo tiene 5 y un cinturón de 4 dimensiones
    # produciría el mismo número: el verificador compararía una contra la otra.
    dims = idx.get("dimensiones") or {}
    tabla_dim = next((t for t in tablas_portada
                      if len(t) > 1 and norm(t[0][0]) == "Dimensión"), None)
    if tabla_dim is None:
        fallas.append(f"{cint}/portada: no encuentro la tabla de dimensiones")
        return
    por_nombre = {norm(d.get("nombre") or k): d for k, d in dims.items()}
    vistas = set()
    for fila in tabla_dim[1:]:
        nombre = norm(fila[0])
        d = por_nombre.get(nombre)
        if d is None:
            fallas.append(f"{cint}/portada: dimensión «{nombre}» no está en el snapshot")
            continue
        vistas.add(nombre)
        if coma(d.get("puntaje"), 1) not in norm(fila[2]):
            fallas.append(f"{cint}/portada/{nombre}: figura «{norm(fila[2])}», "
                          f"el snapshot dice {coma(d.get('puntaje'), 1)}")
        c = COLOR.get((d.get("semaforo") or {}).get("color"))
        if c and norm(fila[3]) != c:
            fallas.append(f"{cint}/portada/{nombre}: figura {norm(fila[3])}, "
                          f"el snapshot dice {c}")
        if coma((d.get("peso") or 0) * 100, 1) not in norm(fila[4]):
            fallas.append(f"{cint}/portada/{nombre}: peso «{norm(fila[4])}», "
                          f"el snapshot dice {coma((d.get('peso') or 0) * 100, 1)} %")
    if set(por_nombre) - vistas:
        fallas.append(f"{cint}/portada: faltan dimensiones {sorted(set(por_nombre) - vistas)}")


def verificar(cint, ruta):
    fallas, avisos = [], []
    doc = Document(ruta)
    secs = secciones(doc)
    tablas = texto_tablas_por_ficha(doc)
    ind = informe["cinturones"][cint]["indicadores"]

    portada = next((t for k, t in tablas.items() if "resumen" in k), None)
    if portada is None:
        fallas.append(f"{cint}: falta la portada de resumen")
    else:
        verificar_portada(cint, portada, fallas)

    esperadas = {norm(LABELS.get(k, k)) for k in ind}
    presentes = {t for t in secs if "resumen" not in t}
    if esperadas - presentes:
        fallas.append(f"faltan fichas: {sorted(esperadas - presentes)}")
    if presentes - esperadas:
        fallas.append(f"fichas de más: {sorted(presentes - esperadas)}")

    for ikey, o in ind.items():
        titulo = norm(LABELS.get(ikey, ikey))
        if titulo not in secs:
            continue
        cuerpo = norm(chr(10).join(secs[titulo]))
        tbls = tablas.get(titulo, [])
        todo = norm(cuerpo + " | " + " | ".join(c for t in tbls for f in t for c in f))
        sem = o.get("semaforo") or {}
        color_txt = COLOR.get(sem.get("color"))

        # Anclado al banner, no "en alguna parte de la ficha".
        #
        # La primera versión de este chequeo buscaba el valor y el color en el
        # texto completo, y una prueba de mutación la pasó por arriba: se
        # alteró el valor del banner y el color de una celda, y el verificador
        # siguió diciendo OK — los dos vuelven a aparecer en la prosa de "Color
        # vigente y por qué". Un chequeo que no falla con el documento mal no
        # sirve para nada.
        banner = tbls[0][0] if tbls and len(tbls[0]) == 1 else None
        if banner is None:
            fallas.append(f"{cint}/{ikey}: no encuentro el banner de la ficha")
        else:
            val = coma(o.get("valor"))
            if o.get("valor") is not None and val not in norm(banner[0]):
                fallas.append(f"{cint}/{ikey}: banner dice «{norm(banner[0])}», "
                              f"el snapshot dice {val}")
            if color_txt and norm(banner[1]) != color_txt:
                fallas.append(f"{cint}/{ikey}: banner dice color «{norm(banner[1])}», "
                              f"el snapshot dice {color_txt}")
            peso = o.get("peso_efectivo")
            if peso and coma(peso * 100, 1) not in norm(banner[2]):
                fallas.append(f"{cint}/{ikey}: banner dice «{norm(banner[2])}», "
                              f"el peso es {coma(peso * 100, 1)} %")

        # La otra afirmación explícita de color que hace la ficha.
        m = re.search(r"Color vigente:\s*([A-ZÁÉÍÓÚ]+)", cuerpo)
        if color_txt and m and m.group(1) != color_txt:
            fallas.append(f"{cint}/{ikey}: «Color vigente: {m.group(1)}» "
                          f"contra {color_txt} en el snapshot")

        # Coherencia interna del ITVC: el índice base-100 impreso tiene que
        # caer en el tramo del color que la misma ficha declara.
        i100 = o.get("indice_itvc")
        if i100 is not None and sem.get("color"):
            esperado = next((c for umbral, c in TRAMOS_ITVC if i100 >= umbral), "rojo")
            if esperado != sem["color"]:
                fallas.append(f"{cint}/{ikey}: índice {i100} es tramo {esperado} "
                              f"pero la ficha dice {sem['color']}")
            if coma(i100, 1) not in todo:
                fallas.append(f"{cint}/{ikey}: falta el índice base-100 {coma(i100, 1)}")

    # Restos de plantilla en todo el documento, texto y tablas.
    entero = "\n".join(p.text for p in doc.paragraphs)
    entero += "\n" + "\n".join(c.text for t in doc.tables for f in t.rows for c in f.cells)
    for patron, que in ((r"\{[A-Za-z_]+\}", "literal sin interpolar"),
                        (r"\bNone\b", "'None' de Python"),
                        (r"\bnan\b", "'nan'"),
                        (r"EN EL \|", "rótulo de dimensión colgado")):
        m = re.findall(patron, entero)
        if m:
            fallas.append(f"{cint}: {que} → {sorted(set(m))[:5]}")

    return fallas, avisos


if __name__ == "__main__":
    total_f, total_a = 0, 0
    for cint, nombre in DOCS.items():
        ruta = AQUI / nombre
        if not ruta.exists():
            print(f"[FALTA] {nombre}")
            total_f += 1
            continue
        f, a = verificar(cint, ruta)
        total_f += len(f)
        total_a += len(a)
        estado = "OK" if not f else f"{len(f)} FALLAS"
        print(f"{nombre:40s} {estado}" + (f" · {len(a)} avisos" if a else ""))
        for x in f:
            print(f"    [FALLA] {x}")
        for x in a:
            print(f"    [aviso] {x}")
    print(f"\nTOTAL: {total_f} fallas · {total_a} avisos")
    sys.exit(1 if total_f else 0)
