"""Estilo de las tablas del .docx y control de los cortes de página.

pandoc arma las tablas sin bordes, sin encabezado distinguido y sin ninguna
de las propiedades de Word que gobiernan qué pasa cuando una tabla cruza el
final de la página. Acá se aplican después, sobre el archivo ya convertido.

Lo que resuelve, en orden de lo que más molesta al leer:

1. Una fila partida al medio por el corte de página (`cantSplit`).
2. Una tabla que sigue en la página siguiente sin encabezado, así que las
   columnas quedan sin nombre (`tblHeader` en la primera fila).
3. Un título al pie de página con su contenido recién en la siguiente
   (`keepNext` en los headings y en el párrafo que los sigue).

Y el estilo: bordes finos, encabezado con fondo, filas alternadas, y las
celdas de color pintadas con la paleta del proyecto — misma que
web/public/dashboard.css, para que la ficha impresa y la página no muestren
dos verdes distintos.
"""
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# Paleta del proyecto (web/public/dashboard.css): fondo suave + texto fuerte.
# El texto va en el tono oscuro y no en el vivo porque el vivo sobre su propio
# soft no llega a contraste legible en impresión.
PALETA = {
    "VERDE":    ("DCFCE7", "14532D"),
    "AMARILLO": ("FEF3C7", "713F12"),
    "NARANJA":  ("FFEDD5", "7C2D12"),
    "ROJO":     ("FEE2E2", "7F1D1D"),
}
GRIS_BORDE = "D9DEDC"
GRIS_HEADER = "1F3A34"     # verde CIGOB oscuro, para la fila de encabezado
GRIS_ALTERNA = "F6F8F7"
GRIS_SECCION = "E8EDEB"    # fila separadora de dimensión
MARCA_DIM = "DIMENSIÓN:"


def _sombrear(celda, hex_color):
    tc = celda._tc.get_or_add_tcPr()
    for viejo in tc.findall(qn("w:shd")):
        tc.remove(viejo)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def _bordes(tabla):
    tbl_pr = tabla._tbl.tblPr
    for viejo in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(viejo)
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")          # 4 octavos de punto = medio punto
        el.set(qn("w:color"), GRIS_BORDE)
        borders.append(el)
    tbl_pr.append(borders)


def _margenes_celda(tabla, pt=4):
    """Aire dentro de la celda. Sin esto el texto toca el borde y la tabla se
    lee apretada por más bordes finos que tenga."""
    tbl_pr = tabla._tbl.tblPr
    for viejo in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(viejo)
    mar = OxmlElement("w:tblCellMar")
    for lado, val in (("top", pt), ("bottom", pt), ("left", pt + 3), ("right", pt + 3)):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:w"), str(int(val * 20)))   # twips
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def _ancho_y_columnas(tabla, ancho_total_twips):
    """La tabla ocupa el ancho completo del texto, con columnas repartidas
    según lo que cada una tiene que mostrar.

    pandoc deja las tablas con el ancho que le parece y alineadas a la
    izquierda: las angostas quedan tiradas contra el margen, con media página
    vacía al lado, y las columnas de una palabra ocupan lo mismo que las de una
    oración. Se calcula el reparto con el largo del texto de cada columna, que
    es la aproximación más simple que da un resultado equilibrado sin tener que
    codificar a mano el diseño de cada tabla.
    """
    filas = tabla.rows
    if not filas:
        return
    n = len(filas[0].cells)
    if n == 0:
        return

    # Dos medidas por columna, y hacen falta las dos:
    #
    #  · el promedio de largo dice cuánto ESPACIO necesita en total, y sirve
    #    para repartir la holgura;
    #  · la palabra más larga dice cuánto ancho necesita como MÍNIMO, porque
    #    una palabra no se parte. Repartir sólo por promedio dejó la columna
    #    de color en 9% y Word puso "VERD / E" y "NARA / NJA" en dos líneas.
    pesos, minimos = [], []
    for j in range(n):
        largos, palabra_larga = [], 4
        for i, fila in enumerate(filas):
            celdas = fila.cells
            if j >= len(celdas):
                continue
            # Las filas separadoras de dimensión abarcan el ancho completo:
            # su texto no dice nada del reparto entre columnas.
            if i > 0 and len(celdas) > 1 and celdas[0].text.strip() == celdas[-1].text.strip():
                continue
            t = celdas[j].text.strip()
            largos.append(len(t) * (0.4 if i == 0 else 1.0))
            # El espacio de no separación NO es un corte posible: "14,0 %" es
            # una sola unidad para Word. `split()` lo trata como whitespace y
            # medía "14,0", así que la columna de peso salía angosta y Word
            # terminaba partiendo el par igual, que es justo lo que el nbsp
            # estaba puesto para evitar.
            for palabra in t.replace(chr(0xA0), chr(0x2013)).split():
                palabra_larga = max(palabra_larga, len(palabra))
        pesos.append(max(sum(largos) / max(len(largos), 1), 3.0))
        # A 9 pt un carácter mide ~5 pt (100 twips); + los márgenes de celda.
        minimos.append(palabra_larga * 100 + 300)

    # Si los mínimos no entran, no hay reparto que salve la tabla: se deja que
    # Word acomode como pueda antes que forzar columnas ilegibles.
    if sum(minimos) >= ancho_total_twips:
        return

    # Reparto: primero el mínimo de cada una, y la holgura por proporción del
    # espacio que cada columna necesita.
    holgura = ancho_total_twips - sum(minimos)
    total_peso = sum(pesos)
    anchos = [m + int(holgura * p / total_peso) for m, p in zip(minimos, pesos)]
    anchos[-1] += ancho_total_twips - sum(anchos)   # el redondeo va a la última

    tbl = tabla._tbl
    tbl_pr = tbl.tblPr
    for etiqueta in ("w:tblW", "w:tblLayout", "w:jc"):
        for viejo in tbl_pr.findall(qn(etiqueta)):
            tbl_pr.remove(viejo)
    w_el = OxmlElement("w:tblW")
    w_el.set(qn("w:w"), str(int(ancho_total_twips)))
    w_el.set(qn("w:type"), "dxa")
    tbl_pr.append(w_el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")     # sin esto Word recalcula y los ignora
    tbl_pr.append(layout)


    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for a in anchos:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(a))
        grid.append(gc)
    tbl.insert(list(tbl).index(tbl_pr) + 1, grid)

    for fila in filas:
        celdas = fila.cells
        for j, celda in enumerate(celdas):
            if j >= len(anchos):
                continue
            tc_pr = celda._tc.get_or_add_tcPr()
            for viejo in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(viejo)
            tw = OxmlElement("w:tcW")
            tw.set(qn("w:w"), str(anchos[j]))
            tw.set(qn("w:type"), "dxa")
            tc_pr.append(tw)


def _aire_alrededor(doc, pt_antes=8, pt_despues=10):
    """Separa las tablas del texto que las rodea.

    Una tabla en Word no tiene espacio propio: pegada al párrafo de arriba y al
    de abajo se lee como si fuera parte de la oración. El aire se pone en los
    párrafos vecinos, que es donde Word lo entiende."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    hijos = list(doc.element.body.iterchildren())
    for i, hijo in enumerate(hijos):
        if not hijo.tag.endswith("}tbl"):
            continue
        if i > 0 and hijos[i - 1].tag.endswith("}p"):
            Paragraph(hijos[i - 1], doc).paragraph_format.space_after = Pt(pt_antes)
        if i + 1 < len(hijos) and hijos[i + 1].tag.endswith("}p"):
            Paragraph(hijos[i + 1], doc).paragraph_format.space_before = Pt(pt_despues)


def _no_partir(fila):
    tr_pr = fila._tr.get_or_add_trPr()
    if not tr_pr.findall(qn("w:cantSplit")):
        tr_pr.append(OxmlElement("w:cantSplit"))


def _repetir_encabezado(fila):
    tr_pr = fila._tr.get_or_add_trPr()
    if not tr_pr.findall(qn("w:tblHeader")):
        tr_pr.append(OxmlElement("w:tblHeader"))


def _keep_next(parrafo):
    p_pr = parrafo._p.get_or_add_pPr()
    if not p_pr.findall(qn("w:keepNext")):
        p_pr.append(OxmlElement("w:keepNext"))


def _combinar_fila(fila):
    """La fila separadora de dimensión ocupa el ancho completo."""
    celdas = fila.cells
    if len(celdas) > 1:
        celdas[0].merge(celdas[-1])


def _texto(celda):
    return celda.text.strip()


def _tiene_encabezado(tabla):
    """Si la primera fila es encabezado de columnas o ya es un dato.

    No se puede preguntar por `tblHeader`: pandoc no lo escribe en ninguna
    tabla, así que a esta altura todas llegan iguales. Y cuando el markdown
    declara el encabezado vacío (`| | | |`, como el banner "Hoy:" y la tabla
    de Identificación), pandoc directamente no emite fila de encabezado — la
    primera fila que se ve es un dato. Pintarla como encabezado deja el valor
    del mes en blanco sobre verde oscuro, que fue exactamente lo que pasó.

    Las dos señales vienen de cómo está armado el markdown de las fichas:
      · una tabla de una sola fila es un banner, nunca tiene encabezado;
      · en esta anatomía las etiquetas de una tabla de datos van en negrita
        ("**IDENTIFICADOR TÉCNICO**") y los nombres de columna de una tabla
        con encabezado, no.
    """
    if len(tabla.rows) == 1:
        return False
    return not any(r.bold for c in tabla.rows[0].cells
                   for p in c.paragraphs for r in p.runs)


# Sello para no aplicar el estilo dos veces sobre el mismo archivo.
#
# Hace falta porque este script NO es idempotente: la primera pasada quita el
# prefijo MARCA_DIM de las filas separadoras, así que la segunda ya no las
# reconoce y las trata como filas comunes — les pisa el sombreado de sección
# con el de fila alternada. Pasó de verdad: tres documentos re-estilados
# perdieron el separador de dimensión y ningún chequeo de contenido lo vio,
# porque el texto seguía intacto.
#
# El pipeline correcto es siempre pandoc → estilar sobre un .docx recién
# convertido. Si hay que re-estilar, se vuelve a convertir primero.
SELLO = "fichas-cigob: estilado"


def estilar(path):
    doc = Document(path)
    if SELLO in (doc.core_properties.comments or ""):
        raise SystemExit(
            f"{path}: ya tiene el estilo aplicado. Volvé a convertirlo con "
            f"pandoc antes de estilarlo — una segunda pasada rompe las filas "
            f"separadoras de dimensión.")
    n_color, n_seccion = 0, 0

    # Ancho útil de la página: es el que tienen que ocupar las tablas.
    #
    # `cigob_reference.docx` no declara el tamaño de página —lo hereda del
    # default de Word— así que `page_width` viene en None y hay que suponerlo.
    # A4 es lo que usa el template en la práctica, verificado sobre el PDF.
    A4_TWIPS = 11906
    EMU_POR_TWIP = 635
    sec = doc.sections[0]
    ancho_pagina = (sec.page_width / EMU_POR_TWIP) if sec.page_width else A4_TWIPS
    margen = ((sec.left_margin or 0) + (sec.right_margin or 0)) / EMU_POR_TWIP
    ancho_texto = int(ancho_pagina - margen)

    for tabla in doc.tables:
        tabla.autofit = False
        _bordes(tabla)
        _margenes_celda(tabla)
        _ancho_y_columnas(tabla, ancho_texto)
        con_encabezado = _tiene_encabezado(tabla)

        for i, fila in enumerate(tabla.rows):
            _no_partir(fila)
            es_seccion = _texto(fila.cells[0]).startswith(MARCA_DIM)

            if i == 0 and con_encabezado:
                _repetir_encabezado(fila)
                for celda in fila.cells:
                    _sombrear(celda, GRIS_HEADER)
                    for p in celda.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.color.rgb = None
                            r.font.size = Pt(9)
                            # El blanco va por XML: font.color.rgb no alcanza
                            # cuando el estilo del template ya fija un color.
                            rPr = r._r.get_or_add_rPr()
                            for v in rPr.findall(qn("w:color")):
                                rPr.remove(v)
                            col = OxmlElement("w:color")
                            col.set(qn("w:val"), "FFFFFF")
                            rPr.append(col)
                continue

            if es_seccion:
                n_seccion += 1
                _combinar_fila(fila)
                for celda in fila.cells:
                    _sombrear(celda, GRIS_SECCION)
                for p in fila.cells[0].paragraphs:
                    p.text = p.text.replace(MARCA_DIM, "").strip()
                    # El separador se pega a la fila que abre: si no, cae al pie
                    # de página anunciando una dimensión cuyos indicadores están
                    # recién en la página siguiente. Pasó con "Vulnerabilidad
                    # financiera" en vida cotidiana.
                    _keep_next(p)
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
                continue

            if con_encabezado and i % 2 == 0:
                for celda in fila.cells:
                    _sombrear(celda, GRIS_ALTERNA)

            for celda in fila.cells:
                fondo_texto = PALETA.get(_texto(celda))
                if fondo_texto:
                    n_color += 1
                    fondo, tinta = fondo_texto
                    _sombrear(celda, fondo)
                    for p in celda.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(9)
                            rPr = r._r.get_or_add_rPr()
                            for v in rPr.findall(qn("w:color")):
                                rPr.remove(v)
                            col = OxmlElement("w:color")
                            col.set(qn("w:val"), tinta)
                            rPr.append(col)
                else:
                    for p in celda.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

    # Un título nunca se queda solo al pie: se pega al párrafo que lo sigue.
    #
    # Solo al título. Marcar además el párrafo siguiente encadena keepNext —
    # título + primer párrafo + segundo párrafo tienen que entrar juntos— y
    # cuando no entran, Word empuja el bloque completo a la página siguiente:
    # media página en blanco después de "Definición", verificado en el PDF.
    n_keep = 0
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            _keep_next(p)
            n_keep += 1

    _aire_alrededor(doc)

    props = doc.core_properties
    props.comments = ((props.comments + " · ") if props.comments else "") + SELLO
    doc.save(path)
    return {"tablas": len(doc.tables), "celdas_color": n_color,
            "filas_seccion": n_seccion, "titulos_pegados": n_keep}


if __name__ == "__main__":
    for archivo in sys.argv[1:]:
        r = estilar(archivo)
        print(f"{archivo:42s} tablas={r['tablas']:3d} color={r['celdas_color']:3d} "
              f"secciones={r['filas_seccion']:2d} keepNext={r['titulos_pegados']:3d}")
